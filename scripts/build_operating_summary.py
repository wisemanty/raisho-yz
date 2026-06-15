#!/usr/bin/env python3
"""Build Markdown and Word meeting summaries from a RAISHO weekly workbook."""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Any

import pandas as pd

from build_weekly_tables import categorize_product, effective_mask, find_col, money, read_table


def value_from_checks(checks: pd.DataFrame, key: str, default: Any = "") -> Any:
    if checks.empty or "检查项" not in checks.columns or "结果" not in checks.columns:
        return default
    rows = checks[checks["检查项"].astype(str) == key]
    if rows.empty:
        return default
    return rows.iloc[0]["结果"]


def format_money(value: Any) -> str:
    try:
        return f"{float(value):,.2f}"
    except (TypeError, ValueError):
        return "0.00"


def safe_text(value: Any) -> str:
    if pd.isna(value):
        return ""
    text = str(value).strip()
    return "" if text.lower() in {"nan", "none", "null"} else text


def read_sheet(workbook: Path, sheet_name: str) -> pd.DataFrame:
    try:
        return pd.read_excel(workbook, sheet_name=sheet_name)
    except Exception:
        return pd.DataFrame()


def load_effective_detail(detail_path: Path | None) -> dict[str, Any]:
    if not detail_path or not detail_path.exists():
        return {}

    raw = read_table(detail_path)
    raw.columns = [str(c).strip() for c in raw.columns]
    uid_col = find_col(raw.columns, ["yz_open_id", "yzOpenId", "有赞客户ID"], required=True)
    order_col = find_col(raw.columns, ["订单号", "订单编号"], required=True)
    status_col = find_col(raw.columns, ["订单状态", "交易状态", "状态"])
    pay_time_col = find_col(raw.columns, ["支付时间", "付款时间"])
    order_time_col = find_col(raw.columns, ["下单时间", "订单创建时间", "创建时间"])
    product_col = find_col(raw.columns, ["商品名称", "商品", "商品标题"], required=True)
    amount_col = find_col(raw.columns, ["商品实收金额", "订单实收", "实收金额", "商品销售金额", "订单金额"], required=True)
    order_amount_col = find_col(raw.columns, ["订单实收", "实收金额", "订单金额"]) or amount_col
    qty_col = find_col(raw.columns, ["商品销售件数", "数量", "件数"])
    time_col = pay_time_col or order_time_col

    df = raw.copy()
    df["_uid"] = df[uid_col].map(safe_text)
    df = df[df["_uid"] != ""].copy()
    df["_order"] = df[order_col].map(safe_text)
    df["_product"] = df[product_col].map(safe_text)
    df["_category"] = df["_product"].map(categorize_product)
    df["_amount"] = money(df[amount_col])
    df["_order_amount"] = money(df[order_amount_col])
    df["_qty"] = money(df[qty_col]) if qty_col else 1
    df["_time"] = pd.to_datetime(df[time_col], errors="coerce") if time_col else pd.NaT
    df["_effective"] = effective_mask(df, status_col)
    eff = df[df["_effective"]].copy()
    return {
        "raw": raw,
        "eff": eff,
        "time_col": time_col or "",
        "status_col": status_col or "",
        "source": str(detail_path),
    }


def summarize_eff(eff: pd.DataFrame, raw_rows: int = 0, time_col: str = "") -> tuple[dict[str, Any], pd.DataFrame]:
    if eff.empty:
        metrics = {
            "原始明细行数": raw_rows,
            "有效明细行数": 0,
            "去重用户数": 0,
            "有效订单数": 0,
            "累计实付": 0.0,
            "客单价": 0.0,
            "最早支付时间": "",
            "最新支付时间": "",
            "时间字段": time_col,
        }
        return metrics, pd.DataFrame(columns=["品类", "客户数", "订单数", "件数", "商品实收"])

    orders = (
        eff.groupby("_order", dropna=False)
        .agg(订单实收=("_order_amount", "max"), yz_open_id=("_uid", "first"), 支付时间=("_time", "min"))
        .reset_index()
    )
    paid = float(orders["订单实收"].sum()) if not orders.empty else 0.0
    order_count = int(orders["_order"].nunique()) if not orders.empty else 0
    metrics = {
        "原始明细行数": raw_rows,
        "有效明细行数": len(eff),
        "去重用户数": int(eff["_uid"].nunique()),
        "有效订单数": order_count,
        "累计实付": paid,
        "客单价": paid / order_count if order_count else 0,
        "最早支付时间": eff["_time"].min(),
        "最新支付时间": eff["_time"].max(),
        "时间字段": time_col,
    }
    product = (
        eff.groupby("_category")
        .agg(
            客户数=("_uid", "nunique"),
            订单数=("_order", "nunique"),
            件数=("_qty", "sum"),
            商品实收=("_amount", "sum"),
        )
        .reset_index()
        .rename(columns={"_category": "品类"})
        .sort_values("商品实收", ascending=False)
    )
    product["商品实收"] = product["商品实收"].round(2)
    return metrics, product


def summarize_detail(detail_path: Path | None) -> tuple[dict[str, Any], pd.DataFrame]:
    loaded = load_effective_detail(detail_path)
    if not loaded:
        return {}, pd.DataFrame()
    return summarize_eff(loaded["eff"], len(loaded["raw"]), loaded["time_col"])


def parse_date_bounds(date_range: str, metrics: dict[str, Any]) -> tuple[pd.Timestamp | None, pd.Timestamp | None]:
    dates = re.findall(r"\d{4}-\d{2}-\d{2}", date_range or "")
    if len(dates) >= 2:
        start = pd.Timestamp(dates[0])
        end = pd.Timestamp(dates[1]) + pd.Timedelta(days=1) - pd.Timedelta(seconds=1)
        return start, end
    start_raw = metrics.get("最早支付时间")
    end_raw = metrics.get("最新支付时间")
    if start_raw == "" or end_raw == "":
        return None, None
    return pd.Timestamp(start_raw).normalize(), pd.Timestamp(end_raw).normalize() + pd.Timedelta(days=1) - pd.Timedelta(seconds=1)


def summarize_period(all_detail: Path | None, start: pd.Timestamp | None, end: pd.Timestamp | None) -> tuple[dict[str, Any], pd.DataFrame]:
    if not all_detail or start is None or end is None:
        return {}, pd.DataFrame()
    loaded = load_effective_detail(all_detail)
    if not loaded:
        return {}, pd.DataFrame()
    eff = loaded["eff"]
    period_eff = eff[(eff["_time"] >= start) & (eff["_time"] <= end)].copy()
    return summarize_eff(period_eff, len(period_eff), loaded["time_col"])


def cumulative_metrics(all_detail: Path | None) -> dict[str, Any]:
    metrics, _ = summarize_detail(all_detail)
    return metrics


def compute_new_old_users(users: pd.DataFrame, all_detail: Path | None, start: pd.Timestamp | None) -> tuple[Any, Any]:
    if users.empty or "yz_open_id" not in users.columns or not all_detail or start is None:
        return "需全量历史对照", "需全量历史对照"
    loaded = load_effective_detail(all_detail)
    if not loaded:
        return "需全量历史对照", "需全量历史对照"
    orders = (
        loaded["eff"].groupby("_uid", dropna=False)
        .agg(first_time=("_time", "min"))
        .reset_index()
    )
    first_times = dict(zip(orders["_uid"], orders["first_time"]))
    weekly_uids = [safe_text(v) for v in users["yz_open_id"].tolist()]
    new_count = 0
    old_count = 0
    for uid in weekly_uids:
        first = first_times.get(uid)
        if pd.isna(first):
            continue
        if pd.Timestamp(first) >= start:
            new_count += 1
        else:
            old_count += 1
    return new_count, old_count


def table_to_md(df: pd.DataFrame, columns: list[str], max_rows: int = 12) -> str:
    if df.empty:
        return "暂无数据"
    view = df[[c for c in columns if c in df.columns]].head(max_rows).copy()
    if view.empty:
        return "暂无数据"
    headers = [str(c) for c in view.columns]
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    for _, row in view.iterrows():
        cells = [safe_text(row[col]).replace("|", "/") for col in view.columns]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def period_words(start: pd.Timestamp | None, end: pd.Timestamp | None) -> dict[str, str]:
    days = None
    if start is not None and end is not None:
        days = (end.normalize() - start.normalize()).days + 1
    if days is not None and days > 7:
        return {
            "title": "来处经营总结",
            "current": "本期",
            "previous": "上期",
            "next": "下一阶段",
            "current_change": "本期变化",
            "current_review": "本期经营动作复盘",
            "previous_action": "上期计划动作",
            "current_handle": "本期处理",
            "current_behavior": "本期/历史行为",
        }
    return {
        "title": "来处周经营总结",
        "current": "本周",
        "previous": "上周",
        "next": "下周",
        "current_change": "本周变化",
        "current_review": "本周经营动作复盘",
        "previous_action": "上周计划动作",
        "current_handle": "本周处理",
        "current_behavior": "本周/历史行为",
    }


def adapt_period_text(text: Any, words: dict[str, str]) -> str:
    value = safe_text(text)
    if words.get("current") == "本周":
        return value
    return (
        value
        .replace("本周给", f"{words['next']}给")
        .replace("本周内", "会后7天内")
        .replace("本周", words["current"])
        .replace("下周", words["next"])
        .replace("上周", words["previous"])
    )


def delta_text(current: Any, previous: Any, unit: str = "", previous_label: str = "上周") -> str:
    try:
        cur = float(current)
        prev = float(previous)
    except (TypeError, ValueError):
        return f"暂无{previous_label}对比"
    diff = cur - prev
    if prev == 0:
        return f"较{previous_label}{'+' if diff >= 0 else ''}{diff:,.2f}{unit}"
    return f"较{previous_label}{'+' if diff >= 0 else ''}{diff:,.2f}{unit}（{diff / prev:+.1%}）"


def product_role(category: str) -> str:
    if category == "玉乃光白标":
        return "成交主力 / 信任入口 / 复购承接"
    if category == "玉乃光黑标":
        return "高信任 / 高客单 / 预售维护"
    if category.startswith("今治"):
        return "礼赠 / 转介绍验证 / 生活方式破冰"
    if category == "AddElm":
        return "配饰补充 / 低客单搭配 / 兴趣测试"
    return "观察类目"


def product_issue(category: str, row: pd.Series) -> str:
    if category == "玉乃光黑标":
        return "样本少，重点看到货后反馈、复购和转介绍。"
    if category == "玉乃光白标":
        return "成交基础已有，需要验证是否能升级黑标或形成稳定复购。"
    if category.startswith("今治"):
        return "仍需验证礼赠和转介绍，不宜只按销量判断。"
    if category == "AddElm":
        return "需要确认是顺手搭配还是可独立经营的品类。"
    return "先保留观察，不做强推判断。"


def product_action(category: str) -> str:
    if category == "玉乃光黑标":
        return "5/31 后做一对一体验回访，记录口感、场景和二次需求。"
    if category == "玉乃光白标":
        return "回访饮用场景，筛选可升级黑标或可进老客群的人。"
    if category.startswith("今治"):
        return "到货后问手感和使用场景，收集可转发反馈。"
    if category == "AddElm":
        return "观察是否适合和酒/生活品做组合，不单独加大资源。"
    return "继续观察成交来源和客户反馈。"


def build_product_judgement(product: pd.DataFrame) -> pd.DataFrame:
    if product.empty:
        return pd.DataFrame(columns=["商品/类目", "本周成交额", "订单数", "客户数", "当前角色", "当前问题", "下周动作"])
    rows = []
    for _, row in product.iterrows():
        category = safe_text(row.get("品类"))
        rows.append({
            "商品/类目": category,
            "本周成交额": round(float(row.get("商品实收", 0) or 0), 2),
            "订单数": int(row.get("订单数", 0) or 0),
            "客户数": int(row.get("客户数", 0) or 0),
            "当前角色": product_role(category),
            "当前问题": product_issue(category, row),
            "下周动作": product_action(category),
        })
    return pd.DataFrame(rows)


def owner_for_user(row: pd.Series) -> str:
    if safe_text(row.get("优先级")) == "P0":
        return "老板/核心运营"
    distributor = safe_text(row.get("分销员归属"))
    if distributor:
        return distributor
    if "今治" in safe_text(row.get("购买品类")):
        return "运营"
    return "运营"


def pool_type(row: pd.Series) -> str:
    priority = safe_text(row.get("优先级"))
    action = safe_text(row.get("下一步动作"))
    group = safe_text(row.get("是否进群"))
    distributor = safe_text(row.get("分销员归属"))
    if priority == "P0" or "一对一" in action:
        return "必须一对一客户"
    if "到货" in action or "体验" in action or "回访" in action:
        return "收货/体验回访客户"
    if group and "暂不" not in group:
        return "可邀请进群客户"
    if distributor:
        return "可由分销员跟进客户"
    return "内容触达客户"


def build_customer_pool(users: pd.DataFrame) -> pd.DataFrame:
    if users.empty:
        return pd.DataFrame(columns=["经营池", "客户", "本周/历史行为", "判断", "下一步动作", "负责人"])
    rows = []
    view = users.sort_values(["优先级", "累计实付"], ascending=[True, False]).head(18)
    for _, row in view.iterrows():
        rows.append({
            "经营池": pool_type(row),
            "客户": safe_text(row.get("客户显示")),
            "本周/历史行为": f"{safe_text(row.get('购买品类'))}；{int(row.get('有效订单数', 0) or 0)}单；累计{format_money(row.get('累计实付', 0))}",
            "判断": safe_text(row.get("用户标签")),
            "下一步动作": safe_text(row.get("下一步动作")),
            "负责人": owner_for_user(row),
        })
    return pd.DataFrame(rows)


def distributor_layer(row: pd.Series) -> str:
    judgement = safe_text(row.get("质量判断"))
    paid = float(row.get("累计实付", 0) or 0)
    customers = int(row.get("带来客户数", 0) or 0)
    if "重点培养" in judgement:
        return "重点培养"
    if paid > 0 and customers >= 2:
        return "稳定观察"
    if paid > 0:
        return "轻培养"
    return "暂不投入"


def build_distributor_rank(distributors: pd.DataFrame) -> pd.DataFrame:
    if distributors.empty:
        return pd.DataFrame(columns=["排名", "分销员", "本周成交额", "有效订单数", "成交客户数", "客单价", "复购客户数", "复购率", "高客单客户数", "黑标客户数", "分层", "下周动作"])
    view = distributors.sort_values(["累计实付", "带来客户数"], ascending=[False, False]).copy()
    rows = []
    for i, (_, row) in enumerate(view.iterrows(), start=1):
        rows.append({
            "排名": i,
            "分销员": safe_text(row.get("分销员")),
            "本周成交额": format_money(row.get("累计实付", 0)),
            "有效订单数": int(row.get("有效订单数", 0) or 0),
            "成交客户数": int(row.get("带来客户数", 0) or 0),
            "客单价": format_money(row.get("客单价", 0)),
            "复购客户数": int(row.get("复购客户数", 0) or 0),
            "复购率": safe_text(row.get("复购率")),
            "高客单客户数": int(row.get("高客单客户数", 0) or 0),
            "黑标客户数": int(row.get("黑标客户数", 0) or 0),
            "分层": distributor_layer(row),
            "下周动作": safe_text(row.get("培养动作")),
        })
    return pd.DataFrame(rows)


def build_distributor_actions(distributor_rank: pd.DataFrame) -> pd.DataFrame:
    if distributor_rank.empty:
        return pd.DataFrame(columns=["分销员", "分层", "下周动作"])
    return distributor_rank[["分销员", "分层", "下周动作"]].copy()


def top_name(df: pd.DataFrame, col: str, fallback: str = "暂无") -> str:
    if df.empty or col not in df.columns:
        return fallback
    return safe_text(df.iloc[0][col]) or fallback


def build_change_lines(
    metrics: dict[str, Any],
    prev_metrics: dict[str, Any],
    product: pd.DataFrame,
    distributors: pd.DataFrame,
    words: dict[str, str],
) -> list[str]:
    top_product = top_name(product, "品类")
    top_dist = top_name(distributors, "分销员")
    compare_text = (
        delta_text(metrics.get("累计实付", 0), prev_metrics.get("累计实付", 0), "元", words["previous"])
        if prev_metrics
        else f"暂无{words['previous']}对比"
    )
    lines = [
        f"成交额 {format_money(metrics.get('累计实付', 0))}，{compare_text}。",
        f"有效订单 {int(metrics.get('有效订单数', 0) or 0)} 单，成交客户 {int(metrics.get('去重用户数', 0) or 0)} 人。",
        f"{words['current']}成交额最高的商品/类目是 {top_product}。",
        f"{words['current']}分销侧最值得复盘的是 {top_dist}。",
    ]
    return lines


def build_action_plan(product_judgement: pd.DataFrame, customer_pool: pd.DataFrame, distributor_rank: pd.DataFrame, words: dict[str, str]) -> pd.DataFrame:
    rows = [
        {"负责人": "老板/核心运营", "动作": "处理 P0 和黑标/高客单客户一对一", "对象": "客户经营池里的必须一对一客户", "截止时间": "会后7天内", "验证指标": "是否获得反馈、复购意向或转介绍线索"},
        {"负责人": "运营", "动作": "完成白标、今治、黑标客户回访分组", "对象": "收货/体验回访客户", "截止时间": "会后7天内", "验证指标": "完成回访人数和可用反馈条数"},
        {"负责人": "运营", "动作": "给重点分销员发送具体商品素材", "对象": "分销员排行里的重点培养/轻培养对象", "截止时间": "会后3天内", "验证指标": "分销员是否转发、咨询、带来新客"},
    ]
    if not product_judgement.empty:
        rows.append({"负责人": "内容", "动作": f"围绕{words['current']}成交主力做一版真实内容", "对象": safe_text(product_judgement.iloc[0].get("商品/类目")), "截止时间": "会后5天内", "验证指标": "是否带来咨询、收藏、进群或转发"})
    if not distributor_rank.empty:
        rows.append({"负责人": "分销员", "动作": "按客户名单轻触达并反馈结果", "对象": f"{words['current']}有成交分销员", "截止时间": "会后7天内", "验证指标": "反馈客户数、二次咨询数、成交数"})
    return pd.DataFrame(rows)


def build_review_table(words: dict[str, str]) -> pd.DataFrame:
    return pd.DataFrame([
        {words["previous_action"]: f"{words['previous']}动作完成记录", "是否完成": "待补入", "结果": "当前 skill 尚未接入动作台账", words["current_handle"]: f"后续接入{words['previous']}行动清单后自动复盘"},
    ])


def build_validation_questions() -> list[str]:
    return [
        "白标客户能不能继续升级到黑标？",
        "黑标客户到货后是否愿意反馈、复购或转介绍？",
        "今治是否真的适合礼赠和转介绍？",
        "哪些分销员能连续带来客户，而不是只偶发成交？",
        "老客能不能被会员、社群或新品优先权承接？",
    ]


def build_summary_text(
    workbook: Path,
    detail_path: Path | None,
    week_label: str,
    date_range: str,
    all_detail_path: Path | None = None,
) -> tuple[str, dict[str, Any]]:
    users = read_sheet(workbook, "01用户分层表")
    distributors = read_sheet(workbook, "03分销员质量表")
    checks = read_sheet(workbook, "07数据质量检查")
    metrics, product = summarize_detail(detail_path)

    if not metrics:
        metrics = {
            "原始明细行数": value_from_checks(checks, "原始行数", 0),
            "有效明细行数": value_from_checks(checks, "有效明细行数", 0),
            "去重用户数": value_from_checks(checks, "去重用户数", 0),
            "有效订单数": value_from_checks(checks, "去重订单数", 0),
            "累计实付": float(users["累计实付"].sum()) if "累计实付" in users.columns else 0,
            "客单价": 0,
            "最早支付时间": "",
            "最新支付时间": "",
            "时间字段": value_from_checks(checks, "使用时间字段", ""),
        }
    repurchase_note = value_from_checks(
        checks,
        "复购统计口径",
        "复购客户数/复购率按购买会话数计算；同一用户同一自然日多单合并为1次购买会话。",
    )

    start_ts, end_ts = parse_date_bounds(date_range, metrics)
    words = period_words(start_ts, end_ts)
    all_detail_for_compare = all_detail_path or detail_path
    cumulative = cumulative_metrics(all_detail_for_compare)
    if all_detail_path and start_ts is not None:
        days = (end_ts.normalize() - start_ts.normalize()).days + 1 if end_ts is not None else 7
        prev_start = start_ts - pd.Timedelta(days=days)
        prev_end = start_ts - pd.Timedelta(seconds=1)
        prev_metrics, _ = summarize_period(all_detail_path, prev_start, prev_end)
    else:
        prev_metrics = {}
    new_customers, old_customers = compute_new_old_users(users, all_detail_path, start_ts)

    product_judgement = build_product_judgement(product).rename(columns={
        "本周成交额": f"{words['current']}成交额",
        "下周动作": f"{words['next']}动作",
    })
    customer_pool = build_customer_pool(users).rename(columns={
        "本周/历史行为": words["current_behavior"],
    })
    raw_distributor_rank = build_distributor_rank(distributors)
    distributor_actions = build_distributor_actions(raw_distributor_rank).rename(columns={
        "下周动作": f"{words['next']}动作",
    })
    distributor_rank = raw_distributor_rank.rename(columns={
        "本周成交额": f"{words['current']}成交额",
        "下周动作": f"{words['next']}动作",
    })
    for df in (distributor_rank, distributor_actions):
        action_col = f"{words['next']}动作"
        if action_col in df.columns:
            df[action_col] = df[action_col].map(lambda value: adapt_period_text(value, words))
    action_review = build_review_table(words)
    action_plan = build_action_plan(product_judgement, customer_pool, distributor_rank, words)
    change_lines = build_change_lines(metrics, prev_metrics, product, distributors, words)
    questions = build_validation_questions()

    top_dist_name = top_name(distributor_rank, "分销员")
    top_product_name = top_name(product, "品类")
    black_amount = 0.0
    if not product.empty and "品类" in product.columns:
        black_amount = float(product.loc[product["品类"].astype(str) == "玉乃光黑标", "商品实收"].sum())
    signal = f"黑标贡献 {format_money(black_amount)}，仍是最强高信任信号" if black_amount > 0 else f"{top_product_name} 是{words['current']}主要成交信号"

    dashboard = pd.DataFrame([
        {"指标": f"{words['current']}成交额", words["current"]: format_money(metrics.get("累计实付", 0)), "累计/对比": f"累计 {format_money(cumulative.get('累计实付', 0))}"},
        {"指标": "有效订单数", words["current"]: int(metrics.get("有效订单数", 0) or 0), "累计/对比": int(cumulative.get("有效订单数", 0) or 0)},
        {"指标": "成交客户数", words["current"]: int(metrics.get("去重用户数", 0) or 0), "累计/对比": int(cumulative.get("去重用户数", 0) or 0)},
        {"指标": "新客/老客", words["current"]: f"新客 {new_customers} / 老客 {old_customers}", "累计/对比": "按 yz_open_id 与历史首购时间判断"},
        {"指标": "客单价", words["current"]: format_money(metrics.get("客单价", 0)), "累计/对比": format_money(cumulative.get("客单价", 0))},
        {"指标": f"{words['current']}对比", words["current"]: delta_text(metrics.get("累计实付", 0), prev_metrics.get("累计实付", 0), "元", words["previous"]) if prev_metrics else f"暂无{words['previous']}对比", "累计/对比": ""},
    ])

    current_amount_col = f"{words['current']}成交额"
    next_action_col = f"{words['next']}动作"

    md = f"""# {words['title']} {week_label}

## 1. {words['current']}一句话结论

{words['current']}成交额 {format_money(metrics.get('累计实付', 0))}，有效订单 {int(metrics.get('有效订单数', 0) or 0)} 单，成交客户 {int(metrics.get('去重用户数', 0) or 0)} 人。{signal}；分销侧以 {top_dist_name} 最值得优先复盘。{words['next']}重点是把高信任客户、白标客户和重点分销员拆成具体跟进动作。

## 2. 核心数据看板

{table_to_md(dashboard, ['指标', words['current'], '累计/对比'], 20)}

## 3. {words['current_change']}

{chr(10).join(f'- {line}' for line in change_lines)}

## 4. 商品经营判断

{table_to_md(product_judgement, ['商品/类目', current_amount_col, '订单数', '客户数', '当前角色', '当前问题', next_action_col], 10)}

## 5. 客户经营池

{table_to_md(customer_pool, ['经营池', '客户', words['current_behavior'], '判断', '下一步动作', '负责人'], 18)}

## 6. 分销员排行与经营判断

### 6.1 分销员销售排行

{table_to_md(distributor_rank, ['排名', '分销员', current_amount_col, '有效订单数', '成交客户数', '客单价', '复购客户数', '复购率', '高客单客户数', '黑标客户数', '分层'], 20)}

### 6.2 分销员{words['next']}动作

{table_to_md(distributor_actions, ['分销员', '分层', next_action_col], 20)}

## 7. {words['current_review']}

{table_to_md(action_review, [words['previous_action'], '是否完成', '结果', words['current_handle']], 10)}

## 8. {words['next']}行动清单

{table_to_md(action_plan, ['负责人', '动作', '对象', '截止时间', '验证指标'], 10)}

## 9. {words['next']}要验证的问题

{chr(10).join(f'- {question}' for question in questions)}

## 10. 附录：数据口径

- 分析区间：{date_range or week_label}
- 数据来源：有赞自助取数 `来处订单商品明细_yz_open_id`
- 分析时间字段：{metrics.get('时间字段', '')}
- 实际最早支付时间：{metrics.get('最早支付时间', '')}
- 实际最新支付时间：{metrics.get('最新支付时间', '')}
- 原始明细行数：{metrics.get('原始明细行数', 0)}
- 有效明细行数：{metrics.get('有效明细行数', 0)}
- 用户去重口径：按 `yz_open_id`
- 有效订单口径：排除关闭、取消、未支付、待付款、全额退款等明确无效状态
- 复购口径：{repurchase_note}
- 分销员口径：当前按分销员昵称聚合；如分销员改名或重名，需要接入分销员 ID / 手机号
- 原始数据文件：`{detail_path or ''}`
"""
    context = {
        "metrics": metrics,
        "cumulative": cumulative,
        "dashboard": dashboard,
        "change_lines": change_lines,
        "product_judgement": product_judgement,
        "customer_pool": customer_pool,
        "distributor_rank": distributor_rank,
        "distributor_actions": distributor_actions,
        "action_review": action_review,
        "action_plan": action_plan,
        "questions": questions,
        "repurchase_note": repurchase_note,
        "detail_path": str(detail_path or ""),
        "words": words,
        "one_sentence": f"{words['current']}成交额 {format_money(metrics.get('累计实付', 0))}，有效订单 {int(metrics.get('有效订单数', 0) or 0)} 单，成交客户 {int(metrics.get('去重用户数', 0) or 0)} 人。{signal}；分销侧以 {top_dist_name} 最值得优先复盘。",
    }
    return md, context


def add_dataframe_table(doc, df: pd.DataFrame, columns: list[str], max_rows: int = 12) -> None:
    from docx.shared import Pt

    view = df[[c for c in columns if c in df.columns]].head(max_rows).copy()
    if view.empty:
        doc.add_paragraph("暂无数据")
        return
    table = doc.add_table(rows=1, cols=len(view.columns))
    table.style = "Table Grid"
    for i, col in enumerate(view.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in view.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(view.columns):
            cells[i].text = safe_text(row[col])
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8)


def add_bullets(doc, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_docx(
    out_docx: Path,
    week_label: str,
    date_range: str,
    context: dict[str, Any],
) -> None:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(20)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 1"].font.color.rgb = RGBColor(11, 92, 126)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(11.5)

    title = doc.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    words = context.get("words", period_words(None, None))
    current_amount_col = f"{words['current']}成交额"
    next_action_col = f"{words['next']}动作"

    title.add_run(words["title"])
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(date_range or week_label).italic = True

    doc.add_heading(f"1. {words['current']}一句话结论", level=1)
    doc.add_paragraph(context["one_sentence"])

    doc.add_heading("2. 核心数据看板", level=1)
    add_dataframe_table(doc, context["dashboard"], ["指标", words["current"], "累计/对比"], 20)

    doc.add_heading(f"3. {words['current_change']}", level=1)
    add_bullets(doc, context["change_lines"])

    doc.add_heading("4. 商品经营判断", level=1)
    add_dataframe_table(doc, context["product_judgement"], ["商品/类目", current_amount_col, "订单数", "客户数", "当前角色", "当前问题", next_action_col], 10)

    doc.add_heading("5. 客户经营池", level=1)
    add_dataframe_table(doc, context["customer_pool"], ["经营池", "客户", words["current_behavior"], "判断", "下一步动作", "负责人"], 18)

    doc.add_heading("6. 分销员排行与经营判断", level=1)
    doc.add_heading("6.1 分销员销售排行", level=2)
    add_dataframe_table(doc, context["distributor_rank"], ["排名", "分销员", current_amount_col, "有效订单数", "成交客户数", "客单价", "复购客户数", "复购率", "高客单客户数", "黑标客户数", "分层"], 20)
    doc.add_heading(f"6.2 分销员{words['next']}动作", level=2)
    add_dataframe_table(doc, context["distributor_actions"], ["分销员", "分层", next_action_col], 20)

    doc.add_heading(f"7. {words['current_review']}", level=1)
    add_dataframe_table(doc, context["action_review"], [words["previous_action"], "是否完成", "结果", words["current_handle"]], 10)

    doc.add_heading(f"8. {words['next']}行动清单", level=1)
    add_dataframe_table(doc, context["action_plan"], ["负责人", "动作", "对象", "截止时间", "验证指标"], 10)

    doc.add_heading(f"9. {words['next']}要验证的问题", level=1)
    add_bullets(doc, context["questions"])

    doc.add_heading("10. 附录：数据口径", level=1)
    metrics = context["metrics"]
    appendix = [
        f"分析区间：{date_range or week_label}",
        "数据来源：有赞自助取数 来处订单商品明细_yz_open_id",
        f"分析时间字段：{metrics.get('时间字段', '')}",
        f"实际最早支付时间：{metrics.get('最早支付时间', '')}",
        f"实际最新支付时间：{metrics.get('最新支付时间', '')}",
        f"原始明细行数：{metrics.get('原始明细行数', 0)}",
        f"有效明细行数：{metrics.get('有效明细行数', 0)}",
        "用户去重口径：按 yz_open_id",
        "有效订单口径：排除关闭、取消、未支付、待付款、全额退款等明确无效状态",
        f"复购口径：{context['repurchase_note']}",
        "分销员口径：当前按分销员昵称聚合；如分销员改名或重名，需要接入分销员 ID / 手机号",
        f"原始数据文件：{context['detail_path']}",
    ]
    add_bullets(doc, appendix)

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)


def build_summary(
    workbook: Path,
    output_dir: Path,
    week_label: str,
    detail: Path | None = None,
    date_range: str = "",
    all_detail: Path | None = None,
) -> tuple[Path, Path]:
    md, context = build_summary_text(workbook, detail, week_label, date_range, all_detail)
    output_dir.mkdir(parents=True, exist_ok=True)
    md_path = output_dir / f"经营分析总结_{week_label}.md"
    docx_path = output_dir / f"经营分析总结_{week_label}.docx"
    md_path.write_text(md, encoding="utf-8")
    build_docx(docx_path, week_label, date_range, context)
    return md_path, docx_path


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--workbook", required=True)
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--week-label", required=True)
    parser.add_argument("--detail")
    parser.add_argument("--all-detail")
    parser.add_argument("--date-range", default="")
    args = parser.parse_args()

    detail = Path(args.detail).expanduser() if args.detail else None
    all_detail = Path(args.all_detail).expanduser() if args.all_detail else None
    md_path, docx_path = build_summary(
        Path(args.workbook).expanduser(),
        Path(args.output_dir).expanduser(),
        args.week_label,
        detail,
        args.date_range,
        all_detail,
    )
    print(md_path)
    print(docx_path)


if __name__ == "__main__":
    main()
